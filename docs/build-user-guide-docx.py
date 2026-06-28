#!/usr/bin/env python3
"""Build AquaWash-User-Guide.docx from the markdown source (UTF-8)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "AquaWash-User-Guide.md"
OUT_PATH = ROOT / "AquaWash-User-Guide.docx"


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_screenshot_box(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ أضف لقطة الشاشة هنا / Insert screenshot here ]")
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True
    doc.add_paragraph()
    # Reserved space for image
    spacer = doc.add_paragraph()
    spacer.add_run("\n\n\n")
    doc.add_paragraph("_" * 60)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= {"-", ":"} for c in row):
            rows.append(row)
        i += 1
    return rows, i


def add_md_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell_text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            table.rows[ri].cells[ci].text = cell_text


def build() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    # Default font — Word will use system fonts for Arabic
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            if in_code_block:
                i += 1
                box_lines: list[str] = []
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    box_lines.append(lines[i])
                    i += 1
                if any("screenshot" in l.lower() or "لقطة" in l for l in box_lines):
                    add_screenshot_box(doc)
                else:
                    add_paragraph(doc, "\n".join(box_lines))
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 0)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_md_table(doc, rows)
            continue

        if stripped.startswith("> "):
            add_paragraph(doc, stripped[2:], bold=False)
            i += 1
            continue

        if stripped.startswith("- ") or re.match(r"^\d+\.\s", stripped):
            add_paragraph(doc, stripped)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**"):
            add_paragraph(doc, stripped.strip("*"), bold=True)
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            i += 1
            continue

        if stripped:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", clean)
            add_paragraph(doc, clean)

        i += 1

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
